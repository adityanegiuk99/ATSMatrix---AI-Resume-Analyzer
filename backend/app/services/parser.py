import io
import re
from dataclasses import dataclass
from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader

SECTION_PATTERNS = {
    "skills": r"(skills|technologies|toolbox)",
    "education": r"(education|degree|university|college)",
    "experience": r"(experience|work history|employment)",
    "projects": r"(projects|portfolio|case studies)",
}

KNOWN_SKILLS = {
    "react", "next.js", "typescript", "javascript", "python", "fastapi", "node.js", "postgresql",
    "redis", "docker", "kubernetes", "aws", "gcp", "openai", "gemini", "rag", "langchain",
    "chromadb", "pinecone", "sql", "graphql", "ci/cd", "tailwind", "machine learning"
}


@dataclass
class ParsedResume:
    raw_text: str
    candidate_name: str
    skills: list[str]
    education: list[str]
    experience: list[str]
    projects: list[str]
    links: list[str]
    parsing_failures: list[str]

    def as_dict(self) -> dict:
        return {
            "skills": self.skills,
            "education": self.education,
            "experience": self.experience,
            "projects": self.projects,
            "links": self.links,
            "parsing_failures": self.parsing_failures,
        }


async def parse_upload(file: UploadFile) -> ParsedResume:
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Resume exceeds 10 MB upload limit")
    suffix = (file.filename or "").lower().split(".")[-1]
    if suffix == "pdf":
        text = _read_pdf(content)
    elif suffix == "docx":
        text = _read_docx(content)
    else:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX resumes are supported")
    return extract_resume(text)


def _read_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_text = [" | ".join(cell.text for cell in row.cells) for table in doc.tables for row in table.rows]
    return "\n".join(paragraphs + table_text)


def extract_resume(text: str) -> ParsedResume:
    normalized = re.sub(r"\s+", " ", text).strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    candidate_name = lines[0][:80] if lines else "Unknown Candidate"
    lower = normalized.lower()
    skills = sorted({skill for skill in KNOWN_SKILLS if skill in lower})
    links = sorted(set(re.findall(r"https?://[^\s)]+|(?:github|linkedin)\.com/[^\s)]+", lower)))
    failures = detect_formatting_risks(text)
    return ParsedResume(
        raw_text=normalized,
        candidate_name=candidate_name,
        skills=skills,
        education=_section_lines(lines, "education"),
        experience=_section_lines(lines, "experience"),
        projects=_section_lines(lines, "projects"),
        links=links,
        parsing_failures=failures,
    )


def _section_lines(lines: list[str], section: str) -> list[str]:
    pattern = re.compile(SECTION_PATTERNS[section], re.I)
    start = next((idx for idx, line in enumerate(lines) if pattern.search(line)), -1)
    if start == -1:
        return []
    output: list[str] = []
    for line in lines[start + 1:start + 7]:
        if any(re.search(p, line, re.I) for key, p in SECTION_PATTERNS.items() if key != section):
            break
        output.append(line[:180])
    return output


def detect_formatting_risks(text: str) -> list[str]:
    risks: list[str] = []
    if "|" in text or "\t" in text:
        risks.append("Tables or tabular columns may collapse into an unreadable sequence in legacy ATS parsers.")
    if re.search(r"[•▪●◦✓★]", text):
        risks.append("Decorative bullets or icons detected; replace with standard hyphen bullets for maximum parser safety.")
    short_lines = [line for line in text.splitlines() if 1 <= len(line.strip()) <= 2]
    if len(short_lines) > 10:
        risks.append("Many short fragments suggest multi-column formatting that could scramble reading order.")
    if len(text.strip()) < 700:
        risks.append("Extracted text is unusually short; the file may be image-based or ATS-unfriendly.")
    return risks
